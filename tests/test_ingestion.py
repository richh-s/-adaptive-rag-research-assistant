from rag_assistant.ingestion.loaders import load_documents
from rag_assistant.ingestion.splitter import split_documents


def _make_minimal_pdf(text_or_pages: str | list[str]) -> bytes:
    """Hand-build a minimal PDF (one or more pages) with real text content streams, so tests
    can exercise actual text extraction without pulling in a PDF-generation dependency.
    Object numbering: 1=Catalog, 2=Pages, 3=Font, then a (Page, Contents) object pair per
    page."""
    pages = [text_or_pages] if isinstance(text_or_pages, str) else text_or_pages
    font_obj_id = 3
    page_obj_ids = [font_obj_id + 1 + 2 * i for i in range(len(pages))]

    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: (
            f"<< /Type /Pages /Kids [{' '.join(f'{pid} 0 R' for pid in page_obj_ids)}] "
            f"/Count {len(pages)} >>"
        ).encode("latin-1"),
        font_obj_id: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for page_obj_id, text in zip(page_obj_ids, pages):
        contents_obj_id = page_obj_id + 1
        objects[page_obj_id] = (
            f"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 {font_obj_id} 0 R >> >> "
            f"/MediaBox [0 0 612 792] /Contents {contents_obj_id} 0 R >>"
        ).encode("latin-1")
        content = f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET".encode("latin-1")
        objects[contents_obj_id] = (
            f"<< /Length {len(content)} >>\nstream\n".encode("latin-1") + content + b"\nendstream"
        )

    ordered_ids = sorted(objects)
    pdf = b"%PDF-1.4\n"
    offsets = {}
    for obj_id in ordered_ids:
        offsets[obj_id] = len(pdf)
        pdf += f"{obj_id} 0 obj\n".encode("latin-1") + objects[obj_id] + b"\nendobj\n"
    xref_offset = len(pdf)
    max_id = ordered_ids[-1]
    pdf += f"xref\n0 {max_id + 1}\n".encode("latin-1")
    pdf += b"0000000000 65535 f \n"
    for obj_id in range(1, max_id + 1):
        pdf += f"{offsets[obj_id]:010d} 00000 n \n".encode("latin-1")
    pdf += f"trailer\n<< /Size {max_id + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode(
        "latin-1"
    )
    return pdf


def test_load_documents_reads_all_supported_files(sample_corpus_dir):
    docs = load_documents(sample_corpus_dir)

    assert len(docs) == 2
    sources = {d.metadata["source"] for d in docs}
    assert sources == {"anthropic.md", "mistral.md"}


def test_load_documents_ignores_unsupported_files(sample_corpus_dir):
    (sample_corpus_dir / "notes.json").write_text("{}")

    docs = load_documents(sample_corpus_dir)

    assert all(d.metadata["source"] != "notes.json" for d in docs)


def test_load_documents_extracts_text_from_pdf(sample_corpus_dir):
    (sample_corpus_dir / "cohere.pdf").write_bytes(
        _make_minimal_pdf("Cohere builds enterprise LLMs.")
    )

    docs = load_documents(sample_corpus_dir)

    pdf_docs = [d for d in docs if d.metadata["source"] == "cohere.pdf"]
    assert len(pdf_docs) == 1
    assert "Cohere builds enterprise LLMs." in pdf_docs[0].page_content


def test_load_documents_skips_corrupt_pdf(sample_corpus_dir):
    (sample_corpus_dir / "broken.pdf").write_bytes(b"%PDF-1.4\nnot a real pdf")

    docs = load_documents(sample_corpus_dir)

    assert all(d.metadata["source"] != "broken.pdf" for d in docs)


def test_load_documents_skips_image_only_pdf(sample_corpus_dir):
    (sample_corpus_dir / "scanned.pdf").write_bytes(_make_minimal_pdf(""))

    docs = load_documents(sample_corpus_dir)

    assert all(d.metadata["source"] != "scanned.pdf" for d in docs)


def test_load_documents_tags_pdf_pages_with_page_number(sample_corpus_dir):
    (sample_corpus_dir / "cohere.pdf").write_bytes(
        _make_minimal_pdf("Cohere builds enterprise LLMs.")
    )

    docs = load_documents(sample_corpus_dir)

    pdf_docs = [d for d in docs if d.metadata["source"] == "cohere.pdf"]
    assert pdf_docs[0].metadata["page"] == 1


def test_load_documents_yields_one_document_per_pdf_page(sample_corpus_dir):
    (sample_corpus_dir / "multi.pdf").write_bytes(
        _make_minimal_pdf(["Page one content.", "Page two content."])
    )

    docs = load_documents(sample_corpus_dir)

    pdf_docs = sorted(
        (d for d in docs if d.metadata["source"] == "multi.pdf"), key=lambda d: d.metadata["page"]
    )
    assert [d.metadata["page"] for d in pdf_docs] == [1, 2]
    assert "Page one content." in pdf_docs[0].page_content
    assert "Page two content." in pdf_docs[1].page_content


def test_split_documents_produces_nonempty_chunks(sample_corpus_dir):
    docs = load_documents(sample_corpus_dir)

    chunks = split_documents(docs, chunk_size=50, chunk_overlap=10)

    assert len(chunks) > len(docs)
    assert all(chunk.page_content.strip() for chunk in chunks)
    assert all(chunk.metadata["source"] in {"anthropic.md", "mistral.md"} for chunk in chunks)


def test_load_documents_extracts_text_and_tables_from_docx(sample_corpus_dir):
    import docx

    document = docx.Document()
    document.add_paragraph("Quarterly revenue grew 40 percent.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "EMEA"
    document.save(str(sample_corpus_dir / "report.docx"))

    docs = load_documents(sample_corpus_dir)

    docx_docs = [d for d in docs if d.metadata["source"] == "report.docx"]
    assert len(docx_docs) == 1
    assert "Quarterly revenue grew 40 percent." in docx_docs[0].page_content
    assert "Region | EMEA" in docx_docs[0].page_content


def test_load_documents_extracts_visible_text_from_html(sample_corpus_dir):
    (sample_corpus_dir / "page.html").write_text(
        "<html><head><title>My Page</title><style>body{color:red}</style></head>"
        "<body><nav>menu</nav><p>Visible article text.</p>"
        "<script>alert('x')</script></body></html>"
    )

    docs = load_documents(sample_corpus_dir)

    html_docs = [d for d in docs if d.metadata["source"] == "page.html"]
    assert len(html_docs) == 1
    assert "Visible article text." in html_docs[0].page_content
    assert "alert" not in html_docs[0].page_content
    assert "menu" not in html_docs[0].page_content
    assert html_docs[0].metadata["title"] == "My Page"


def test_load_documents_skips_empty_html(sample_corpus_dir):
    (sample_corpus_dir / "empty.html").write_text("<html><body><script>x()</script></body></html>")

    docs = load_documents(sample_corpus_dir)

    assert all(d.metadata["source"] != "empty.html" for d in docs)


def _png_bytes(width: int = 200, height: int = 150) -> bytes:
    """Minimal solid-color PNG via pymupdf -- no PIL dependency needed."""
    import pymupdf

    pixmap = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, width, height))
    pixmap.clear_with(120)
    return pixmap.tobytes("png")


def _pdf_with_image(path, width: int = 200, height: int = 150, with_text: bool = True) -> None:
    """Builds a real PDF containing an embedded raster image (and optionally text)."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    if with_text:
        page.insert_text((72, 72), "Quarterly report text.")
    page.insert_image(
        pymupdf.Rect(72, 100, 72 + width, 100 + height), stream=_png_bytes(width, height)
    )
    doc.save(str(path))
    doc.close()


def test_pdf_figures_are_described_when_vision_enabled(sample_corpus_dir, monkeypatch):
    monkeypatch.setenv("PDF_VISION", "true")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    from rag_assistant.config import get_settings

    get_settings.cache_clear()
    monkeypatch.setattr(
        "rag_assistant.ingestion.vision.describe_image",
        lambda image_bytes, media_type, prompt: "Bar chart of revenue by region.",
    )

    _pdf_with_image(sample_corpus_dir / "report.pdf")
    docs = load_documents(sample_corpus_dir)

    pdf_docs = [d for d in docs if d.metadata["source"] == "report.pdf"]
    assert len(pdf_docs) == 1
    assert "Quarterly report text." in pdf_docs[0].page_content
    assert "[Figure on page 1: Bar chart of revenue by region.]" in pdf_docs[0].page_content


def test_pdf_tiny_images_are_not_described(sample_corpus_dir, monkeypatch):
    monkeypatch.setenv("PDF_VISION", "true")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    from rag_assistant.config import get_settings

    get_settings.cache_clear()
    calls = []
    monkeypatch.setattr(
        "rag_assistant.ingestion.vision.describe_image",
        lambda image_bytes, media_type, prompt: calls.append(1) or "icon",
    )

    # 40x40 is below MIN_IMAGE_DIMENSION_PX -- logo/icon territory.
    _pdf_with_image(sample_corpus_dir / "logo.pdf", width=40, height=40)
    docs = load_documents(sample_corpus_dir)

    assert calls == []
    logo_docs = [d for d in docs if d.metadata["source"] == "logo.pdf"]
    assert "[Figure" not in logo_docs[0].page_content


def test_scanned_pdf_is_transcribed_when_vision_enabled(sample_corpus_dir, monkeypatch):
    monkeypatch.setenv("PDF_VISION", "true")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    from rag_assistant.config import get_settings

    get_settings.cache_clear()
    from rag_assistant.ingestion import vision as vision_module

    def _fake_describe(image_bytes, media_type, prompt):
        if prompt is vision_module.SCANNED_PAGE_PROMPT:
            return "INVOICE #42 -- Total due: $1,300"
        return "figure description"

    monkeypatch.setattr("rag_assistant.ingestion.vision.describe_image", _fake_describe)

    # Image-only page: no text layer at all, like a scan.
    _pdf_with_image(sample_corpus_dir / "scan.pdf", width=400, height=500, with_text=False)
    docs = load_documents(sample_corpus_dir)

    scan_docs = [d for d in docs if d.metadata["source"] == "scan.pdf"]
    assert len(scan_docs) == 1
    assert "INVOICE #42" in scan_docs[0].page_content


def test_image_only_pdf_still_skipped_when_vision_disabled(sample_corpus_dir):
    # PDF_VISION=false (conftest default): behavior matches the pre-vision loader.
    _pdf_with_image(sample_corpus_dir / "scan2.pdf", with_text=False)
    docs = load_documents(sample_corpus_dir)

    assert all(d.metadata["source"] != "scan2.pdf" for d in docs)
