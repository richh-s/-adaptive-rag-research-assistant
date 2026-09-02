import hashlib
import logging
from dataclasses import dataclass
from pathlib import Path

import pymupdf
import pymupdf4llm
from bs4 import BeautifulSoup
from langchain_core.documents import Document

from rag_assistant.ingestion import vision
from rag_assistant.ingestion.ownership import owner_of_relative_path

SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx", ".html", ".htm"}

# Bumped when a loader's *output* for unchanged bytes changes -- a new PDF extraction path,
# a different HTML boilerplate rule. Recorded in the ingestion manifest alongside the file
# fingerprint, so improving a loader re-indexes affected files instead of leaving the
# collection full of text the old loader produced. Same mechanism as CHUNKING_VERSION.
LOADER_VERSION = 1

logger = logging.getLogger(__name__)

# 2x zoom ≈ 144 dpi -- enough for legible OCR-style transcription without huge payloads.
_SCAN_RENDER_ZOOM = 2.0


def _describe_page_figures(doc: "pymupdf.Document", page_index: int, budget: list[int]) -> str:
    """Vision pass over one page's embedded images: returns "[Figure on page N: ...]" blocks
    to append to the page text, so charts/diagrams/photos become searchable. `budget` is a
    single-element mutable counter shared across the whole PDF (MAX_IMAGES_PER_PDF)."""
    blocks = []
    try:
        images = doc[page_index].get_images(full=True)
    except Exception:
        logger.warning("Could not enumerate images on page %d", page_index + 1, exc_info=True)
        return ""

    for image_info in images:
        if budget[0] <= 0:
            logger.info("Figure-description budget exhausted; remaining images skipped")
            break
        try:
            extracted = doc.extract_image(image_info[0])
        except Exception:
            continue
        if (
            extracted.get("width", 0) < vision.MIN_IMAGE_DIMENSION_PX
            or extracted.get("height", 0) < vision.MIN_IMAGE_DIMENSION_PX
        ):
            continue
        budget[0] -= 1
        description = vision.describe_image(
            extracted["image"], f"image/{extracted.get('ext', 'png')}", vision.FIGURE_PROMPT
        )
        if description:
            blocks.append(f"[Figure on page {page_index + 1}: {description}]")
    return "\n\n".join(blocks)


def _transcribe_scanned_page(doc: "pymupdf.Document", page_index: int) -> str:
    """A page with no text layer is a scan/photo: render it to PNG and let the vision model
    transcribe it -- same mechanism as figure description, no OCR dependency."""
    try:
        pixmap = doc[page_index].get_pixmap(
            matrix=pymupdf.Matrix(_SCAN_RENDER_ZOOM, _SCAN_RENDER_ZOOM)
        )
        png_bytes = pixmap.tobytes("png")
    except Exception:
        logger.warning("Could not render page %d for transcription", page_index + 1, exc_info=True)
        return ""
    return vision.describe_image(png_bytes, "image/png", vision.SCANNED_PAGE_PROMPT) or ""


def _load_docx(path: Path) -> list[Document]:
    """Extract paragraphs and tables from a .docx in document order. Tables are flattened to
    one pipe-joined line per row -- crude, but it keeps cell values adjacent to their row
    context, which is what retrieval actually needs from a table."""
    import docx  # deferred: python-docx import is slow enough to matter at API startup

    try:
        document = docx.Document(str(path))
    except Exception:
        logger.warning("Skipping unreadable DOCX: %s", path.name, exc_info=True)
        return []

    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    if not text.strip():
        logger.warning("No extractable text in DOCX: %s", path.name)
        return []
    return [Document(page_content=text, metadata={"source": path.name})]


def extract_html_text(html: str) -> tuple[str | None, str]:
    """Shared by the .html file loader and URL ingestion: returns (title, visible text).
    Boilerplate tags (nav/script/style/etc.) are dropped wholesale rather than running a
    full readability pass -- good enough for docs/articles, and dependency-free beyond bs4."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(
        ["script", "style", "noscript", "template", "nav", "footer", "header", "aside"]
    ):
        tag.decompose()
    title = soup.title.get_text(strip=True) if soup.title else None
    lines = [line.strip() for line in soup.get_text(separator="\n").splitlines()]
    text = "\n".join(line for line in lines if line)
    return title, text


def _load_html(path: Path) -> list[Document]:
    try:
        title, text = extract_html_text(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        logger.warning("Skipping unreadable HTML: %s", path.name, exc_info=True)
        return []

    if not text.strip():
        logger.warning("No extractable text in HTML: %s", path.name)
        return []
    metadata: dict = {"source": path.name}
    if title:
        metadata["title"] = title
    return [Document(page_content=text, metadata=metadata)]


def _load_pdf(path: Path) -> list[Document]:
    """Extract text natively as Markdown (preserves headings/tables/lists far better than
    raw text extraction) with one Document per page, so citations can point at an exact
    page instead of just the file. A PDF therefore contributes *multiple* Documents sharing
    one `source` -- callers must group by `metadata["source"]` rather than assume one
    Document per file, which held for every format before PDFs were added."""
    try:
        pages = pymupdf4llm.to_markdown(str(path), page_chunks=True)
    except Exception:
        # Corrupt/encrypted/unreadable PDF -- skip rather than crash the whole ingestion run.
        logger.warning("Skipping unreadable PDF: %s", path.name, exc_info=True)
        return []

    # One fitz handle for the vision passes (figure extraction, scanned-page rendering).
    # Opened lazily and only when vision is on -- plain text-only ingestion never pays for it.
    vision_doc: pymupdf.Document | None = None
    if vision.vision_available():
        try:
            vision_doc = pymupdf.open(str(path))
        except Exception:
            logger.warning("Could not reopen %s for vision passes", path.name, exc_info=True)
    figure_budget = [vision.MAX_IMAGES_PER_PDF]

    documents = []
    for page_index, page in enumerate(pages):
        text = (page.get("text") or "").strip()
        page_number = page.get("metadata", {}).get("page_number")

        if vision_doc is not None:
            if not text:
                # No text layer: a scanned/photographed page -- transcribe it.
                text = _transcribe_scanned_page(vision_doc, page_index).strip()
            figures = _describe_page_figures(vision_doc, page_index, figure_budget)
            if figures:
                text = f"{text}\n\n{figures}".strip()

        if not text:
            continue
        documents.append(
            Document(page_content=text, metadata={"source": path.name, "page": page_number})
        )

    if vision_doc is not None:
        vision_doc.close()

    if not documents:
        # Image-only PDF and vision disabled/unavailable -- nothing to index.
        logger.warning("No extractable text in PDF: %s", path.name)
    return documents


@dataclass(frozen=True)
class CorpusFile:
    """One indexable file, identified without parsing it.

    `fingerprint` hashes the raw bytes rather than the parsed text, which is the whole point:
    deciding whether a file needs re-indexing must not require the expensive step that
    re-indexing would perform. Parsing a PDF runs pymupdf4llm and, with PDF_VISION on, a
    vision API call per figure and per scanned page -- so hashing parsed content to decide
    whether to skip parsing pays the entire cost it was meant to avoid.
    """

    path: Path
    source: str
    owner: str
    fingerprint: str


def _fingerprint(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def iter_corpus_files(source_dir: Path, owner: str | None = None) -> list[CorpusFile]:
    """Enumerates indexable files and fingerprints them, without parsing any of them.

    `owner` narrows the scan to one tenant's scope -- their subtree, or just the flat public
    files for the public tenant. `None` scans everything, which is what a CLI `ingest` or a
    container's startup index wants. Scoping matters because ingestion is triggered per
    upload: without it, one tenant adding a 2KB note walks, fingerprints, and re-indexes
    against every other tenant's documents.
    """
    files: list[CorpusFile] = []
    for path in sorted(source_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix not in SUPPORTED_SUFFIXES:
            logger.warning("Skipping unsupported file: %s", path.name)
            continue

        relative = path.relative_to(source_dir)
        file_owner = owner_of_relative_path(relative)
        if owner is not None and file_owner != owner:
            continue
        files.append(
            CorpusFile(
                path=path,
                source=relative.as_posix(),
                owner=file_owner,
                fingerprint=_fingerprint(path),
            )
        )
    return files


def load_corpus_file(corpus_file: CorpusFile) -> list[Document]:
    """Parses one enumerated file. This is the expensive step every caller wants to skip."""
    path = corpus_file.path
    if path.suffix == ".pdf":
        loaded = _load_pdf(path)
    elif path.suffix == ".docx":
        loaded = _load_docx(path)
    elif path.suffix in (".html", ".htm"):
        loaded = _load_html(path)
    else:
        loaded = [Document(page_content=path.read_text(encoding="utf-8"), metadata={})]

    # The per-format loaders set `source` to the bare filename (and PDFs add `page`);
    # overwrite with the relative path and attach the owner in one place here, so a new
    # loader can't forget to do it and quietly publish a tenant's file to everyone.
    for document in loaded:
        document.metadata = {
            **document.metadata,
            "source": corpus_file.source,
            "owner": corpus_file.owner,
        }
    return loaded


def load_documents(source_dir: Path, owner: str | None = None) -> list[Document]:
    """Parse every supported file under source_dir. Convenience wrapper over
    `iter_corpus_files` + `load_corpus_file` for callers that genuinely want everything;
    incremental ingestion uses the two-step form so it can skip parsing unchanged files.

    Walks recursively (see ownership.py): flat files are the shared public corpus, files
    under `_t/<owner>/` belong to that tenant. `source` is the relative POSIX path rather
    than the bare filename, so two tenants can upload files with the same name without one
    silently overwriting the other's manifest entry and chunk IDs -- for flat public files
    the relative path *is* the filename, so nothing about the baseline corpus changes.
    """
    documents: list[Document] = []
    for corpus_file in iter_corpus_files(source_dir, owner=owner):
        documents.extend(load_corpus_file(corpus_file))
    return documents
